import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_table_from_markdown(doc, table_text):
    """마크다운 테이블을 Word 문서에 추가"""
    lines = table_text.strip().split('\n')
    if len(lines) < 3:
        return
    
    # 테이블 헤더 파싱
    header_line = lines[0].strip('|').strip()
    headers = [h.strip() for h in header_line.split('|')]
    
    # 데이터 행 파싱
    rows = []
    for line in lines[2:]:
        if line.strip() and not line.startswith('|---'):
            row = [cell.strip() for cell in line.strip('|').strip().split('|')]
            if len(row) == len(headers):
                rows.append(row)
    
    # Word 테이블 생성
    if rows:
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Light Grid Accent 1'
        
        # 헤더 행
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            # 헤더 스타일
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
        
        # 데이터 행
        for row_data in rows:
            row = table.add_row()
            for i, cell_data in enumerate(row_data):
                row.cells[i].text = cell_data
                for paragraph in row.cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

def markdown_to_docx(markdown_file, output_file):
    """마크다운 파일을 Word 문서로 변환"""
    with open(markdown_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    doc = Document()
    
    # 섹션 변수들
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # H1 제목
        if line.startswith('# '):
            title = line[2:].strip()
            p = doc.add_paragraph(title, style='Heading 1')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(16)
                run.font.bold = True
            i += 1
        
        # H2 제목
        elif line.startswith('## '):
            title = line[3:].strip()
            p = doc.add_paragraph(title, style='Heading 2')
            for run in p.runs:
                run.font.size = Pt(14)
                run.font.bold = True
            i += 1
        
        # H3 제목
        elif line.startswith('### '):
            title = line[4:].strip()
            p = doc.add_paragraph(title, style='Heading 3')
            for run in p.runs:
                run.font.size = Pt(12)
                run.font.bold = True
            i += 1
        
        # 테이블
        elif line.strip().startswith('|'):
            table_lines = [line]
            i += 1
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            
            table_text = '\n'.join(table_lines)
            add_table_from_markdown(doc, table_text)
        
        # 굵은 텍스트 (**text**)
        elif line.strip().startswith('**') and line.strip().endswith('**'):
            p = doc.add_paragraph()
            text = line.strip()[2:-2]
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(11)
            i += 1
        
        # 일반 텍스트
        elif line.strip() and not line.startswith(('|', '#', '---')):
            if line.strip() != '':
                p = doc.add_paragraph(line.strip())
                for run in p.runs:
                    run.font.size = Pt(11)
            i += 1
        
        else:
            i += 1
    
    # 문서 저장
    doc.save(output_file)
    return output_file

# 변환할 파일들
base_path = r'c:\Users\rkfka\Documents\Job\QA-Projects\Testing Project 3 - E-commerce Web App\2. Test Design'
files = [
    'API_Testing_Prioritized_Test_Cases.md',
    'System_Testing_Prioritized_Test_Cases.md',
    'UAT_Prioritized_Test_Cases.md',
    'Risk_Based_Test_Priority_Matrix.md',
    'Priority_Calculation_Examples.md'
]

print("Converting markdown to Word documents...")
for file in files:
    md_path = os.path.join(base_path, file)
    docx_path = os.path.join(base_path, file.replace('.md', '.docx'))
    
    if os.path.exists(md_path):
        markdown_to_docx(md_path, docx_path)
        print(f"✓ {file} -> {file.replace('.md', '.docx')}")
    else:
        print(f"✗ File not found: {file}")

print("\nConversion completed!")
